# sourcestack-api/app/parsing.py
import io
import re
from typing import Optional, Tuple, List
import phonenumbers
from phonenumbers import NumberParseException
import pytesseract
from pdfminer.high_level import extract_text as pdfminer_extract
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine, LTFigure, LTChar
import pypdfium2 as pdfium
from docx import Document

def _extract_pdf_hyperlinks(data: bytes) -> List[str]:
    """Extract hyperlinks/URLs from PDF annotations and text."""
    urls = []
    try:
        from pdfminer.pdfparser import PDFParser
        from pdfminer.pdfdocument import PDFDocument
        from pdfminer.pdfpage import PDFPage
        
        fp = io.BytesIO(data)
        parser = PDFParser(fp)
        doc = PDFDocument(parser)
        
        # Extract URLs from annotations
        for page in PDFPage.create_pages(doc):
            if '/Annots' in page.attrs:
                annots = page.attrs['/Annots']
                for annot in annots:
                    if '/A' in annot:
                        action = annot['/A']
                        if '/URI' in action:
                            uri = action['/URI']
                            if isinstance(uri, str):
                                urls.append(uri)
    except Exception:
        # If hyperlink extraction fails, continue without it
        pass
    
    return urls

def pdf_text_with_ocr_fallback(data: bytes) -> Tuple[str, bool]:
    """
    Extract text from PDF. Use OCR if text extraction yields minimal content.
    Also extracts hyperlinks and appends them to text.
    Returns (text, ocr_used).
    """
    ocr_used = False
    
    # Try pdfminer first
    try:
        text = pdfminer_extract(io.BytesIO(data))
        
        # Extract hyperlinks from PDF and append to text
        hyperlinks = _extract_pdf_hyperlinks(data)
        if hyperlinks:
            # Append hyperlinks to text so they can be extracted
            text += "\n" + "\n".join(hyperlinks)
        
        # If we got very little text (< 50 chars), try OCR
        if len(text.strip()) < 50:
            ocr_used = True
            text = _ocr_pdf(data)
    except Exception:
        # Fallback to OCR on any error
        ocr_used = True
        text = _ocr_pdf(data)
    
    return text, ocr_used

def _ocr_pdf(data: bytes) -> str:
    """Extract text from PDF using OCR."""
    pdf = pdfium.PdfDocument(data)
    text_parts = []
    
    for page_num in range(len(pdf)):
        page = pdf.get_page(page_num)
        # Render page to image
        pil_image = page.render(scale=2.0).to_pil()
        # Extract text via OCR
        page_text = pytesseract.image_to_string(pil_image)
        text_parts.append(page_text)
    
    return "\n".join(text_parts)

def docx_text(data: bytes) -> str:
    """Extract text from DOCX file."""
    doc = Document(io.BytesIO(data))
    paragraphs = [para.text for para in doc.paragraphs]
    return "\n".join(paragraphs)

def extract_email(text: str) -> Optional[str]:
    """Extract email address from text using regex. Handles mailto: links and href attributes."""
    # First check for mailto: links (href="mailto:email@example.com" or mailto:email@example.com)
    mailto_patterns = [
        r'mailto:\s*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})',  # mailto:email@example.com
        r'href=["\']mailto:([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})["\']',  # href="mailto:email@example.com"
    ]
    
    for pattern in mailto_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            return matches[0].lower()
    
    # Check for email near "email" keyword (handles cases like "Email: href='mailto:...'")
    email_keyword_context = re.search(
        r'(?:email|e-mail|mail)[\s:]*.*?(?:href=["\'])?(?:mailto:)?([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})',
        text,
        re.IGNORECASE
    )
    if email_keyword_context:
        return email_keyword_context.group(1).lower()
    
    # Common email pattern (fallback)
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(email_pattern, text)
    if matches:
        return matches[0].lower()
    
    return None

def normalize_phone(text: str) -> Optional[str]:
    """
    Extract and normalize phone number to E.164 format.
    If 10-digit Indian local number, assume +91 prefix.
    """
    # First try direct parsing (handles formatted numbers like +1-555-123-4567)
    try:
        parsed = phonenumbers.parse(text, None)
        if phonenumbers.is_valid_number(parsed):
            return phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.E164)
    except NumberParseException:
        pass
    
    # Remove common separators and normalize whitespace
    cleaned = re.sub(r'[\s\-\(\)\.]', '', text)
    
    # Find sequences of digits (7-15 digits) - remove word boundaries to catch numbers in text
    phone_pattern = r'\d{7,15}'
    matches = re.findall(phone_pattern, cleaned)
    
    for match in matches:
        # If 10 digits, assume Indian (+91)
        if len(match) == 10:
            candidate = f"+91{match}"
        else:
            # Try with + prefix if it looks like a country code
            if len(match) >= 10:
                candidate = f"+{match}"
            else:
                candidate = match
        
        # Try parsing with phonenumbers
        try:
            parsed = phonenumbers.parse(candidate, None)
            if phonenumbers.is_valid_number(parsed):
                return phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.E164)
        except NumberParseException:
            continue
    
    return None

def extract_linkedin(text: str) -> Optional[str]:
    """Extract LinkedIn profile URL or username from text. Handles href attributes and links near 'LinkedIn' keyword."""
    # First check for href attributes with LinkedIn URLs
    href_patterns = [
        r'href=["\'](https?://(?:www\.)?linkedin\.com/in/[a-zA-Z0-9\-]+)["\']',  # href="https://linkedin.com/in/username"
        r'href=["\'](linkedin\.com/in/[a-zA-Z0-9\-]+)["\']',  # href="linkedin.com/in/username"
    ]
    
    for pattern in href_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            url = matches[0]
            if not url.startswith('http'):
                url = f"https://www.{url}"
            return url
    
    # Check for LinkedIn URL near "LinkedIn" keyword (handles cases like "LinkedIn: href='...'")
    linkedin_keyword_context = re.search(
        r'(?:linkedin|linked\s*in)[\s:]*.*?(?:href=["\'])?(https?://(?:www\.)?linkedin\.com/in/[a-zA-Z0-9\-]+)',
        text,
        re.IGNORECASE
    )
    if linkedin_keyword_context:
        return linkedin_keyword_context.group(1)
    
    # Pattern for LinkedIn URLs: linkedin.com/in/username or linkedin.com/profile/view?id=...
    linkedin_patterns = [
        r'https?://(?:www\.)?linkedin\.com/in/([a-zA-Z0-9\-]+)',  # Full URL
        r'linkedin\.com/in/([a-zA-Z0-9\-]+)',  # linkedin.com/in/username
        r'www\.linkedin\.com/in/([a-zA-Z0-9\-]+)',  # www.linkedin.com/in/username
        r'linkedin\.com/profile/view\?id=([a-zA-Z0-9\-]+)',  # Old format
    ]
    
    for pattern in linkedin_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            username = matches[0]
            # Return full URL format
            return f"https://www.linkedin.com/in/{username}"
    
    # Also check for just "linkedin.com/in/" without capturing username
    if re.search(r'linkedin\.com/in/', text, re.IGNORECASE):
        # Try to extract the full URL
        url_match = re.search(r'https?://(?:www\.)?linkedin\.com/in/[a-zA-Z0-9\-]+', text, re.IGNORECASE)
        if url_match:
            return url_match.group(0)
    
    return None

def extract_github(text: str) -> Optional[str]:
    """Extract GitHub profile URL or username from text. Handles href attributes and links near 'GitHub' keyword."""
    # First check for href attributes with GitHub URLs
    href_patterns = [
        r'href=["\'](https?://(?:www\.)?github\.com/[a-zA-Z0-9](?:[a-zA-Z0-9]|-(?=[a-zA-Z0-9])){0,38})["\']',  # href="https://github.com/username"
        r'href=["\'](github\.com/[a-zA-Z0-9](?:[a-zA-Z0-9]|-(?=[a-zA-Z0-9])){0,38})["\']',  # href="github.com/username"
    ]
    
    for pattern in href_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            url = matches[0]
            if not url.startswith('http'):
                url = f"https://{url}"
            return url
    
    # Check for GitHub URL near "GitHub" keyword (handles cases like "GitHub: href='...'")
    github_keyword_context = re.search(
        r'(?:github|git\s*hub)[\s:]*.*?(?:href=["\'])?(https?://(?:www\.)?github\.com/[a-zA-Z0-9](?:[a-zA-Z0-9]|-(?=[a-zA-Z0-9])){0,38})',
        text,
        re.IGNORECASE
    )
    if github_keyword_context:
        return github_keyword_context.group(1)
    
    # Pattern for GitHub URLs: github.com/username
    github_patterns = [
        r'https?://(?:www\.)?github\.com/([a-zA-Z0-9](?:[a-zA-Z0-9]|-(?=[a-zA-Z0-9])){0,38})',  # Full URL
        r'github\.com/([a-zA-Z0-9](?:[a-zA-Z0-9]|-(?=[a-zA-Z0-9])){0,38})',  # github.com/username (GitHub username rules)
        r'www\.github\.com/([a-zA-Z0-9](?:[a-zA-Z0-9]|-(?=[a-zA-Z0-9])){0,38})',  # www.github.com/username
    ]
    
    for pattern in github_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            username = matches[0]
            # Return full URL format
            return f"https://github.com/{username}"
    
    # Also check for just "github.com/" without capturing username
    if re.search(r'github\.com/', text, re.IGNORECASE):
        # Try to extract the full URL
        url_match = re.search(r'https?://(?:www\.)?github\.com/[a-zA-Z0-9](?:[a-zA-Z0-9]|-(?=[a-zA-Z0-9])){0,38}', text, re.IGNORECASE)
        if url_match:
            return url_match.group(0)
    
    return None

def extract_fields(text: str) -> Tuple[Optional[str], Optional[str], Optional[str], Optional[str]]:
    """Extract email, phone, LinkedIn, and GitHub from text."""
    email = extract_email(text)
    phone = normalize_phone(text)
    linkedin = extract_linkedin(text)
    github = extract_github(text)
    return email, phone, linkedin, github

def guess_name(text: str) -> Optional[str]:
    """
    Heuristically guess name from resume text.
    Looks at first ~30 lines and lines above email/phone/contact keywords.
    """
    lines = text.split('\n')
    # Look at first 30 lines
    candidate_lines = lines[:30]
    
    # Also look for lines above email/phone
    email_phone_keywords = ['email', 'phone', 'contact', 'mobile', 'tel']
    for i, line in enumerate(lines[:50]):
        line_lower = line.lower()
        if any(kw in line_lower for kw in email_phone_keywords) and i > 0:
            candidate_lines.append(lines[i-1])
    
    # Filter out empty lines and common non-name patterns
    name_pattern = re.compile(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+$')
    for line in candidate_lines:
        line = line.strip()
        if not line:
            continue
        # Skip if looks like email, phone, or section header
        if '@' in line or re.match(r'^\+?\d', line) or len(line) > 50:
            continue
        # Check if it matches name pattern (2-4 words, capitalized)
        words = line.split()
        if 2 <= len(words) <= 4:
            if all(word[0].isupper() if word else False for word in words):
                return line
    
    return None

def score_confidence(name: Optional[str], email: Optional[str], 
                    phone: Optional[str], linkedin: Optional[str],
                    github: Optional[str], ocr_used: bool) -> float:
    """
    Calculate confidence score.
    Weights: email 0.4, phone 0.25, name 0.15, linkedin 0.1, github 0.05, +0.05 if not OCR.
    """
    score = 0.0
    if email:
        score += 0.4
    if phone:
        score += 0.25
    if name:
        score += 0.15
    if linkedin:
        score += 0.1
    if github:
        score += 0.05
    if not ocr_used:
        score += 0.05
    
    return min(score, 1.0)

def parse_resume_bytes(filename: str, data: bytes) -> Tuple[dict, list[str], bool]:
    """
    Parse resume file and extract fields.
    Returns (parsed_dict, errors_list, ocr_used).
    """
    errors = []
    name = None
    email = None
    phone = None
    linkedin = None
    github = None
    ocr_used = False
    text = ""
    
    try:
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            text, ocr_used = pdf_text_with_ocr_fallback(data)
        elif filename_lower.endswith('.docx'):
            text = docx_text(data)
        else:
            errors.append(f"Unsupported file type: {filename}")
            return {
                "name": None,
                "email": None,
                "phone": None,
                "linkedin": None,
                "github": None,
                "confidence": 0.0
            }, errors, False
        
        # Extract fields
        email, phone, linkedin, github = extract_fields(text)
        name = guess_name(text)
        
        confidence = score_confidence(name, email, phone, linkedin, github, ocr_used)
        
        return {
            "name": name,
            "email": email,
            "phone": phone,
            "linkedin": linkedin,
            "github": github,
            "confidence": confidence
        }, errors, ocr_used
    
    except Exception as e:
        errors.append(f"Parse error: {str(e)}")
        return {
            "name": None,
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None,
            "confidence": 0.0
        }, errors, False

